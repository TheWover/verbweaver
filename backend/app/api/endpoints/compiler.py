from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional, Dict
import os
import tempfile
from pathlib import Path
import subprocess
import aiofiles
import markdown
from weasyprint import HTML
from docx import Document
import ebooklib
from ebooklib import epub

from app.database import get_db
from app.models import User, Project
from app.core.security import get_current_user
from app.core.config import settings
from pydantic import BaseModel

router = APIRouter()


class CompileOptions(BaseModel):
    includeMetadata: bool = True
    includeToc: bool = True
    includeIndex: bool = False
    includeBibliography: bool = False
    pageSize: str = "A4"
    fontSize: str = "medium"
    margins: str = "normal"
    lineSpacing: str = "1.5"


class CompileRequest(BaseModel):
    format: str
    options: CompileOptions
    nodes: Optional[List[str]] = None  # Specific nodes to include


def get_project_path(project_id: int) -> Path:
    """Get the filesystem path for a project."""
    return Path(settings.PROJECTS_DIR) / f"project_{project_id}"


async def collect_markdown_files(project_path: Path, nodes: Optional[List[str]] = None) -> str:
    """Collect all markdown files in the project and combine them."""
    content = []
    
    # If specific nodes are requested, only include those
    if nodes:
        for node_id in nodes:
            file_path = project_path / f"{node_id}.md"
            if file_path.exists():
                async with aiofiles.open(file_path, 'r') as f:
                    content.append(await f.read())
    else:
        # Include all markdown files
        for file_path in sorted(project_path.rglob("*.md")):
            async with aiofiles.open(file_path, 'r') as f:
                content.append(await f.read())
    
    return "\n\n---\n\n".join(content)


def apply_formatting_options(content: str, options: CompileOptions) -> str:
    """Apply formatting options to the content."""
    # Add CSS for formatting
    css = """
    <style>
    body {
        font-family: 'Times New Roman', serif;
    """
    
    # Font size
    if options.fontSize == "small":
        css += "font-size: 12pt;"
    elif options.fontSize == "large":
        css += "font-size: 16pt;"
    else:
        css += "font-size: 14pt;"
    
    # Line spacing
    if options.lineSpacing == "single":
        css += "line-height: 1.0;"
    elif options.lineSpacing == "double":
        css += "line-height: 2.0;"
    else:
        css += "line-height: 1.5;"
    
    # Margins
    if options.margins == "narrow":
        css += "margin: 1cm;"
    elif options.margins == "wide":
        css += "margin: 3cm;"
    else:
        css += "margin: 2cm;"
    
    css += """
    }
    h1 { page-break-before: always; }
    </style>
    """
    
    # Convert markdown to HTML
    html_content = markdown.markdown(content, extensions=['extra', 'toc'])
    
    # Add table of contents if requested
    if options.includeToc:
        toc = '<div class="toc"><h2>Table of Contents</h2>[TOC]</div>'
        html_content = toc + html_content
    
    return css + html_content


async def compile_to_pdf(content: str, options: CompileOptions, output_path: Path):
    """Compile content to PDF."""
    html_content = apply_formatting_options(content, options)
    
    # Use WeasyPrint to convert HTML to PDF
    HTML(string=html_content).write_pdf(output_path)


async def compile_to_docx(content: str, options: CompileOptions, output_path: Path):
    """Compile content to DOCX."""
    doc = Document()
    
    # Add content
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    
    doc.save(output_path)


async def compile_to_epub(content: str, options: CompileOptions, output_path: Path):
    """Compile content to EPUB."""
    book = epub.EpubBook()
    
    # Set metadata
    book.set_identifier('verbweaver-doc')
    book.set_title('Document')
    book.set_language('en')
    
    # Create chapter
    c1 = epub.EpubHtml(title='Document', file_name='chap_01.xhtml', lang='en')
    c1.content = apply_formatting_options(content, options)
    
    # Add chapter
    book.add_item(c1)
    
    # Add navigation
    book.toc = (epub.Link('chap_01.xhtml', 'Document', 'doc'),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Create spine
    book.spine = ['nav', c1]
    
    # Write to file
    epub.write_epub(output_path, book, {})


async def compile_to_html(content: str, options: CompileOptions, output_path: Path):
    """Compile content to HTML."""
    html_content = apply_formatting_options(content, options)
    
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Document</title>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    async with aiofiles.open(output_path, 'w') as f:
        await f.write(full_html)


@router.post("/projects/{project_id}/compile")
async def compile_document(
    project_id: int,
    compile_request: CompileRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Compile project documents to various formats."""
    # Check project access
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found"
        )
    
    project_path = get_project_path(project_id)
    
    if not project_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project files not found"
        )
    
    # Collect content
    content = await collect_markdown_files(project_path, compile_request.nodes)
    
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No content to compile"
        )
    
    # Create temporary file for output
    output_dir = Path(tempfile.mkdtemp())
    filename = f"document.{compile_request.format}"
    output_path = output_dir / filename
    
    try:
        # Compile based on format
        if compile_request.format == "pdf":
            await compile_to_pdf(content, compile_request.options, output_path)
        elif compile_request.format == "docx":
            await compile_to_docx(content, compile_request.options, output_path)
        elif compile_request.format == "epub":
            await compile_to_epub(content, compile_request.options, output_path)
        elif compile_request.format == "html":
            await compile_to_html(content, compile_request.options, output_path)
        elif compile_request.format in ["odt", "mobi"]:
            # For ODT and MOBI, we'll first convert to HTML then use pandoc
            html_path = output_dir / "temp.html"
            await compile_to_html(content, compile_request.options, html_path)
            
            # Use pandoc for conversion
            subprocess.run([
                "pandoc",
                str(html_path),
                "-o",
                str(output_path)
            ], check=True)
        elif compile_request.format == "markdown":
            # Just save the markdown content
            async with aiofiles.open(output_path, 'w') as f:
                await f.write(content)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported format: {compile_request.format}"
            )
        
        # Return the file
        return FileResponse(
            path=str(output_path),
            media_type='application/octet-stream',
            filename=filename
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Compilation failed: {str(e)}"
        )


@router.get("/projects/{project_id}/compile/formats")
async def get_supported_formats(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get list of supported export formats."""
    return {
        "formats": [
            {"id": "pdf", "name": "PDF", "description": "Portable Document Format"},
            {"id": "docx", "name": "Word", "description": "Microsoft Word Document"},
            {"id": "odt", "name": "OpenDocument", "description": "OpenDocument Text"},
            {"id": "epub", "name": "EPUB", "description": "Electronic Publication"},
            {"id": "mobi", "name": "MOBI", "description": "Kindle Format"},
            {"id": "html", "name": "HTML", "description": "Web Page"},
            {"id": "markdown", "name": "Markdown", "description": "Plain Text with Formatting"}
        ]
    } 